# Pump Line Calculator – SI (with Word export)
# Rezultate mari: NPSH(A), p2 (refulare), p1 (aspiratie)
# Afiseaza ΔH_S / ΔH_D (m), recomanda Q_min / Q_max si exporta raport DOCX.

import math
from io import BytesIO

import pandas as pd
import streamlit as st
from docx import Document

# ---------- Stil "wide" + fara ancore ----------
st.markdown("""
<style>
h1 a, h2 a, h3 a, h4 a, h5 a, h6 a {
    display: none;
}
</style>
""", unsafe_allow_html=True)

g = 9.80665  # m/s²

# ---------- Biblioteca de fittinguri ----------
LIB = pd.DataFrame([
    ["Elbow","Elbow 90° (LR)","K",0.35,"0.3–0.4"],
    ["Elbow","Elbow 45°","K",0.18,"0.15–0.2"],
    ["Tee","Tee – run","K",0.60,"0.5–0.7"],
    ["Tee","Tee – branch","K",1.80,"1.5–2.0"],
    ["Valve","Ball (full bore)","K",0.05,"≈neglijabil"],
    ["Valve","Gate (open)","K",0.15,"0.1–0.2"],
    ["Valve","Globe (open)","K",10.0,"8–12"],
    ["Valve","Butterfly (open)","K",0.70,"0.5–1.0"],
    ["Check","Check (swing)","K",2.00,"1.5–2.5"],
    ["Meter","Debitmetru electromagnetic","K",0.30,"vendor data"],
    ["Strainer","Y-strainer clean","dp",0.05,"bar/buc"],
    ["Strainer","Basket clean","dp",0.10,"bar/buc"],
    ["Filter","Cartridge clean","dp",0.20,"bar/buc"],
    ["Reducer","Reducer gradual","K",0.30,"0.2–0.5"],
    ["Instr.","Probe (LSL/FSL)","K",0.05,"≈neglijabil"],
], columns=["grp","item","tip","val","note"])

# ---------- Helpers ----------
def Re(rho, v, D, mu): 
    return rho * v * D / mu

def f_swamee_jain(Rey, eps, D):
    if Rey <= 0: 
        return 0.0
    return 0.25 / (math.log10(eps/(3.7*D) + 5.74/(Rey**0.9))**2)

def h_major(f, L, D, v): 
    return f * (L/D) * (v*v) / (2*g)

def h_minor(Ksum, v):     
    return Ksum * (v*v) / (2*g)

def dp_pa(rho, h):        # head [m] -> Pa
    return rho * g * h

def pa_to_bar(pa):        # Pa -> bar
    return pa / 1e5

def bar_to_pa(bar):       # bar -> Pa
    return bar * 1e5

def p_atm_from_alt_bar(h_m: float) -> float:
    """Presiune atmosferică standard în bar(abs) la altitudine h [m]."""
    pa = 101325.0 * (1.0 - 2.25577e-5 * h_m) ** 5.25588
    return pa / 1e5

def ui_line(title, Dmm_def, L_def, epsmm_def, key):
    """UI pentru o linie (S sau D). Returneaza D[m], L[m], eps[m], Ksum[-], dpsum_bar[bar]."""
    st.markdown(f"### {title}")
    c1,c2,c3 = st.columns(3)
    with c1: 
        D_mm = st.number_input("D [mm]", 5.0, 2000.0, Dmm_def, 1.0, key=key+"D")
        D = D_mm / 1000.0
    with c2: 
        L = st.number_input("L [m]", 0.0, 20000.0, L_def, 0.5, key=key+"L")
    with c3: 
        eps = st.number_input("ε [mm]", 0.0, 2.0, epsmm_def, 0.005, key=key+"eps") / 1000.0

    st.caption("Alege fittingurile (dp = Δp fix [bar]/buc)")
    items = st.multiselect("Fittings", LIB["item"].tolist(), default=[], key=key+"sel")

    rows = []
    for it in items:
        r = LIB[LIB.item == it].iloc[0]
        q = st.number_input(f"{it} – Qty", 0, 200, 1, key=key+it)
        if q > 0:
            rows.append({"item": it, "tip": r["tip"], "val": r["val"], "qty": q})
    df = pd.DataFrame(rows) if rows else pd.DataFrame(columns=["item","tip","val","qty"])

    # ΣK si ΣΔp(bar)
    Ksum = float((df.query("tip=='K'")["val"] * df.query("tip=='K'")["qty"]).sum()) if not df.empty else 0.0
    dpsum_bar = float((df.query("tip=='dp'")["val"] * df.query("tip=='dp'")["qty"]).sum()) if not df.empty else 0.0

    # campuri suplimentare
    K_extra = st.number_input("K (extra) [-]", 0.0, 500.0, 0.0, 0.1, key=key+"Kextra")
    dp_extra_bar = st.number_input("Δp suplimentar [bar]", 0.0, 5.0, 0.0, 0.01, key=key+"dpExtra")

    Ksum += K_extra
    dpsum_bar += dp_extra_bar

    return D, L, eps, Ksum, dpsum_bar, df

def generate_report(NPSH_a, p1_bar, p2_bar, h_S, h_D, inputs_tbl):
    """Genereaza un raport Word si il intoarce ca BytesIO."""
    doc = Document()

    # Coperta / pagina de completat
    doc.add_heading("Raport calcul pompă", 0)
    doc.add_paragraph(
        "Spațiu liber pentru completare manuală:\n\n"
        "TAG pompă: __________________________\n"
        "Proiect: ____________________________\n"
        "Data: ______________________________\n"
    )
    doc.add_page_break()

    # Rezultate principale
    doc.add_heading("Rezultate principale", level=1)
    doc.add_paragraph(f"NPSH (A): {NPSH_a:.3f} m")
    doc.add_paragraph(f"Presiune în stutul de aspirație (p₁): {p1_bar:.3f} bar(abs)")
    doc.add_paragraph(f"Presiune în stutul de refulare (p₂): {p2_bar:.3f} bar(abs)")
    doc.add_paragraph(f"Pierderea pe conductă de aspirație (ΔH_S): {h_S:.3f} m")
    doc.add_paragraph(f"Pierderea pe conductă de refulare (ΔH_D): {h_D:.3f} m")

    # Date de intrare
    doc.add_heading("Date de intrare", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Mărime", "Valoare", "Unități"
    for _, row in inputs_tbl.iterrows():
        rr = table.add_row().cells
        rr[0].text = str(row["Mărime"])
        # rotunjesc elegant daca e numeric
        val = row["Valoare"]
        try:
            val = f"{float(val):.3f}"
        except Exception:
            val = str(val)
        rr[1].text = val
        rr[2].text = str(row["Unități"])

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ================== Panou stanga (intrari) ==================
st.sidebar.title("Intrări (SI)")

rho = st.sidebar.number_input("ρ [kg/m³]", 100.0, 3000.0, 1000.0, 1.0)
mu  = st.sidebar.number_input("μ [mPa·s]", 0.1, 5000.0, 1.0, 0.1) / 1000.0  # Pa·s
Q   = st.sidebar.number_input("Q [m³/h]", 0.01, 200000.0, 50.0, 0.1) / 3600.0  # m³/s
Pv  = bar_to_pa(st.sidebar.number_input("P_vap [bar(abs)]", min_value=0.0, max_value=15.0,
                                        value=0.023, step=0.001, format="%.3f"))

pos = st.sidebar.radio("Pompa vs suprafața de aspirație",
                       ["Sub nivel (inundată)","Deasupra nivelului"], index=0)
dz_mag = st.sidebar.number_input("Distanța verticală |Δz| [m]", 0.0, 200.0, 2.0, 0.1)
Dz_S = dz_mag if pos == "Sub nivel (inundată)" else -dz_mag

res = st.sidebar.radio("Rezervor de aspirație", ["Atmosferică","Custom"], index=0)
if res == "Atmosferică":
    alt = st.sidebar.number_input("Altitudine [m]", -400.0, 4000.0, 0.0, 10.0)
    Ps_bar = p_atm_from_alt_bar(alt)
    st.sidebar.write(f"P_atm ≈ **{Ps_bar:.3f} bar(abs)**")
else:
    Ps_bar = st.sidebar.number_input("P_s [bar(abs)]", 0.0, 50.0, 1.013, 0.01)
Ps = bar_to_pa(Ps_bar)

Pd_bar = st.sidebar.number_input("P_d (destinație) [bar(abs)]", 0.0, 100.0, 1.013, 0.01)
Pd = bar_to_pa(Pd_bar)

Dz_D = st.sidebar.number_input("Δz_D [m]", -200.0, 500.0, 10.0, 0.1)

# ================== Titlu ==================
st.title("Pump Line Calculator")

# ================== Linii S & D ==================
D_S, L_S, eps_S, K_S, dpsumS_bar, _ = ui_line("S – Aspirație", 100.0, 20.0, 0.015, "S")
D_D, L_D, eps_D, K_D, dpsumD_bar, _ = ui_line("D – Refulare", 80.0, 80.0, 0.045, "D")

# ================== Calcule ==================
A_S, A_D = math.pi * D_S**2 / 4.0, math.pi * D_D**2 / 4.0
v_S, v_D = Q / A_S, Q / A_D

Re_S, Re_D = Re(rho, v_S, D_S, mu), Re(rho, v_D, D_D, mu)
f_S, f_D   = f_swamee_jain(Re_S, eps_S, D_S), f_swamee_jain(Re_D, eps_D, D_D)

h_S = h_major(f_S, L_S, D_S, v_S) + h_minor(K_S, v_S) + bar_to_pa(dpsumS_bar) / (rho * g)  # [m]
h_D = h_major(f_D, L_D, D_D, v_D) + h_minor(K_D, v_D) + bar_to_pa(dpsumD_bar) / (rho * g)  # [m]

p1 = Ps + dp_pa(rho, Dz_S - h_S)     # [Pa]
p2 = Pd + dp_pa(rho, Dz_D + h_D)     # [Pa]

NPSH_a = p1 / (rho * g) + v_S**2 / (2 * g) - Pv / (rho * g)

# Recomandari Q_min / Q_max (dupa viteza pe S)
vmin, vmax = 0.5, 1.5
Qmin, Qmax = vmin * A_S * 3600.0, vmax * A_S * 3600.0  # m³/h

# ================== Rezultate ==================
st.markdown("---")
st.markdown("<div class='big'>Rezultate</div>", unsafe_allow_html=True)
c1,c2,c3 = st.columns(3)
with c1: st.metric("NPSH (A) [m]", f"{NPSH_a:.3f}")
with c2: st.metric("Presiune refulare p₂ [bar(abs)]", f"{pa_to_bar(p2):.3f}")
with c3: st.metric("Presiune stut aspirație p₁ [bar(abs)]", f"{pa_to_bar(p1):.3f}")

st.markdown("#### Pierderi pe conducte (head)")
c4,c5 = st.columns(2)
with c4: st.metric("ΔH_S – Aspirație [m]", f"{h_S:.3f}")
with c5: st.metric("ΔH_D – Refulare [m]", f"{h_D:.3f}")

st.markdown("---")
st.write(f"**Sugestie debit** (D_S={D_S*1000:.0f} mm; v_S recomandat {vmin}…{vmax} m/s): "
         f"Q_min ≈ **{Qmin:.1f} m³/h**, Q_max ≈ **{Qmax:.1f} m³/h**.")

# Tabel intrari pentru raport
tbl = pd.DataFrame({
    "Mărime": ["ρ","μ","Q","D_S","L_S","ε_S","K_S","Δp_S[bar]","D_D","L_D","ε_D","K_D","Δp_D[bar]","P_s","P_d","Δz_S","Δz_D"],
    "Valoare":[rho, mu*1000, Q*3600, D_S*1000, L_S, eps_S*1000, K_S, dpsumS_bar,
               D_D*1000, L_D, eps_D*1000, K_D, dpsumD_bar, Ps_bar, Pd_bar, Dz_S, Dz_D],
    "Unități":["kg/m³","mPa·s","m³/h","mm","m","mm","-","bar","mm","m","mm","-","bar","bar(abs)","bar(abs)","m","m"]
})

# Buton export Word
report_file = generate_report(NPSH_a, pa_to_bar(p1), pa_to_bar(p2), h_S, h_D, tbl)
st.download_button(
    label="📄 Descarcă raport Word",
    data=report_file,
    file_name="raport_pompa.docx",
    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)

